import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


path_local = Path.cwd()
# Путь к папке с документами и изображением
folder_path = path_local
image_path = os.path.join(folder_path, "name.png")  # Укажите имя файла изображения

# Параметры вставки изображения
distance_from_bottom = Inches(1)  # Расстояние от нижнего края страницы (1 дюйм)

# Перебор всех файлов .docx в папке
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        doc_path = os.path.join(folder_path, filename)
        doc = Document(doc_path)

        # Проверка, что в документе есть хотя бы 2 страницы
        if len(doc.sections) > 0:
            # Выбор второго раздела (второй страницы)
            section = doc.sections[0]

            # Добавление изображения на второй лист
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(2))  # Укажите нужную ширину изображения

            # Установка положения изображения
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            # Добавление пустого абзаца для создания отступа от нижнего края
            for _ in range(10):  # Настройте количество пустых абзацев для нужного отступа
                doc.add_paragraph()

            # Сохранение изменённого документа
            doc.save(doc_path)
            print(f"Изображение добавлено в документ: {filename}")
        else:
            print(f"В документе {filename} недостаточно страниц для вставки изображения.")
