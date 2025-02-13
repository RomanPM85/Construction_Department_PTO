import os
from docx import Document


def replace_text_in_docx(doc_path, old_text, new_text):
    # Открываем документ
    doc = Document(doc_path)

    # Проходим по всем параграфам и заменяем текст
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)

    # Проходим по всем таблицам и заменяем текст
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

    # Сохраняем изменения
    doc.save(doc_path)


def replace_text_in_folder(folder_path, old_text, new_text):
    # Рекурсивно обходим все файлы и папки
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(root, file)
                print(f"Обрабатываем файл: {file_path}")
                replace_text_in_docx(file_path, old_text, new_text)


if __name__ == "__main__":
    # Укажите путь к текущей папке
    current_folder = os.getcwd()

    # Укажите текст для замены
    old__text = "старый текст"
    new__text = "новый текст"

    # Запускаем замену
    replace_text_in_folder(current_folder, old__text, new__text)

