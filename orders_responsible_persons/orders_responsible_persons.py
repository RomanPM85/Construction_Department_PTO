from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

def create_order_doc(filename, title, city, date, order_number, content_lines, signatures):
    doc = Document()

    # Параметры страницы: формат А4, портрет, поля 2 см
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4 высота
    section.page_width = Inches(8.27)    # A4 ширина
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Стиль основного текста
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Добавляем под заголовком две строки с названием компании и ИНН + телефон
    p_company = doc.add_paragraph()
    p_company.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_company = p_company.add_run('ООО "Строй - Монтаж"')
    run_company.font.size = Pt(12)
    run_company.bold = True

    p_info = doc.add_paragraph()
    p_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_info = p_info.add_run('ИНН 1234567890, тел.: +7 (495) 123-45-67')
    run_info.font.size = Pt(12)

    doc.add_paragraph()

    # Город и дата на одной строке, с противоположных сторон
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = True

    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.add_run(f"г. {city}")

    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p_right.add_run(f"«{date}»")

    doc.add_paragraph()

    # Номер приказа — увеличенный шрифт и по центру
    p_order = doc.add_paragraph()
    p_order.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_order = p_order.add_run(f"Приказ № {order_number}")
    run_order.bold = True
    run_order.font.size = Pt(16)

    # Заголовок приказа — чуть меньше шрифт, по центру, под номером приказа
    p_title = doc.add_paragraph()
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(14)

    doc.add_paragraph()

    # Включаем информацию об объекте в нужную строку content_lines
    for i, line in enumerate(content_lines):
        if line.startswith("В целях обеспечения"):
            content_lines[i] = line.rstrip('.') + " на объекте 'Здание Физкультурно-оздоровительного комплекса, расположенное по адресу: МО, Одинцовский городской округ, п. Барвиха, КИЗ Яблоневый Сад'."
            break

    # Основной текст приказа
    for line in content_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.first_line_indent = Cm(0.63)
        p.paragraph_format.space_after = Pt(3)
        p.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.add_paragraph()

    # Добавляем строку "С приказом ознакомлен:"
    p_ack = doc.add_paragraph()
    p_ack.paragraph_format.space_before = Pt(18)
    p_ack.add_run("С приказом ознакомлен:")

    # Подписи
    p_sign1 = doc.add_paragraph()
    p_sign1.paragraph_format.space_before = Pt(18)
    p_sign1.add_run(f"Генеральный директор: _______________ /подпись/\n")
    p_sign1.add_run(f"{signatures.get('director', '')}")

    p_sign2 = doc.add_paragraph()
    p_sign2.paragraph_format.space_before = Pt(12)
    p_sign2.add_run(f"Ответственный: _______________ /подпись/\n")
    p_sign2.add_run(f"{signatures.get('responsible', '')}")

    doc.save(filename)



if __name__ == "__main__":
    city = "Москва"
    date = "03.07.2025"
    director_name = "Иванов И.И."

    orders = [
        {
            "filename": "Приказ_по_организации_работ_на_высоте.docx",
            "title": "О назначении ответственного за организацию работ на высоте на ",
            "order_number": "20",
            "content": [
                "В целях обеспечения безопасности при выполнении работ на высоте,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за организацию работ на высоте следующего сотрудника:",
                "   ФИО: Козлов Алексей Дмитриевич",
                "   Должность: Инженер по охране труда",
                "2. Ответственному лицу обеспечить:",
                "   - организацию безопасного выполнения высотных работ;",
                "   - контроль за соблюдением требований охраны труда и техники безопасности;",
                "   - проведение инструктажей и обучение персонала;",
                "   - контроль исправности средств индивидуальной защиты и страховочных систем.",
                "3. Контроль за исполнением настоящего приказа возложить на начальника отдела охраны труда."
            ],
            "signatures": {"director": director_name, "responsible": "Козлов А.Д."}
        },
        {
            "filename": "Приказ_о_проведении_инструктажа_по_охране_труда.docx",
            "title": "О проведении инструктажа по охране труда",
            "order_number": "21",
            "content": [
                "В целях повышения уровня безопасности труда ,",
                "ПРИКАЗЫВАЮ:",
                "1. Организовать проведение вводного и первичного инструктажа по охране труда для всех новых работников.",
                "2. Ответственным за проведение инструктажей назначить следующего сотрудника:",
                "   ФИО: Морозова Елена Сергеевна",
                "   Должность: Специалист по охране труда",
                "3. Обеспечить ведение журналов инструктажей и своевременное обновление знаний работников.",
                "4. Контроль за исполнением настоящего приказа возложить на начальника отдела охраны труда."
            ],
            "signatures": {"director": director_name, "responsible": "Морозова Е.С."}
        },
        {
            "filename": "Приказ_о_допуске_к_высотным_работам.docx",
            "title": "О допуске работников к выполнению высотных работ",
            "order_number": "22",
            "content": [
                "В целях обеспечения безопасности при выполнении высотных работ,",
                "ПРИКАЗЫВАЮ:",
                "1. Допустить к выполнению высотных работ следующих работников:",
                "   - Иванов Сергей Петрович",
                "   - Смирнова Ольга Викторовна",
                "2. Работники должны иметь соответствующую подготовку и пройти инструктаж по охране труда.",
                "3. Ответственным за контроль допуска назначить:",
                "   ФИО: Козлов Алексей Дмитриевич",
                "   Должность: Инженер по охране труда",
                "4. Контроль за исполнением настоящего приказа возложить на начальника строительного участка."
            ],
            "signatures": {"director": director_name, "responsible": "Козлов А.Д."}
        },
        {
            "filename": "Приказ_по_пожарной_безопасности.docx",
            "title": "О назначении ответственного за пожарную безопасность",
            "order_number": "15",
            "content": [
                "В целях обеспечения пожарной безопасности на строительной площадке и предупреждения пожаров,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за пожарную безопасность следующего сотрудника:",
                "   ФИО: Иванов Иван Иванович",
                "   Должность: Инженер по охране труда",
                "2. Ответственному лицу обеспечить:",
                "   - контроль за соблюдением требований пожарной безопасности;",
                "   - организацию проведения инструктажей по пожарной безопасности для работников;",
                "   - контроль исправности средств пожаротушения и эвакуационных путей;",
                "   - своевременное информирование о возможных нарушениях.",
                "3. Контроль за исполнением настоящего приказа возложить на начальника отдела охраны труда.",
                "Основание: договор подряда № 123 от «01» 07 2025 г."
            ],
            "signatures": {"director": director_name, "responsible": "Иванов И.И."}
        },
        {
            "filename": "Приказ_по_средствам_пожаротушения.docx",
            "title": "О назначении ответственного за эксплуатацию средств пожаротушения",
            "order_number": "16",
            "content": [
                "В целях обеспечения готовности средств пожаротушения на объекте,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за эксплуатацию и техническое обслуживание средств пожаротушения:",
                "   ФИО: Петров Петр Петрович",
                "   Должность: Техник по охране труда",
                "2. Ответственному лицу обеспечить:",
                "   - регулярную проверку и техническое обслуживание огнетушителей и другого оборудования;",
                "   - своевременное пополнение и замену средств пожаротушения;",
                "   - ведение учета и отчетности по средствам пожаротушения.",
                "3. Контроль за исполнением настоящего приказа возложить на начальника отдела охраны труда."
            ],
            "signatures": {"director": director_name, "responsible": "Петров П.П."}
        },
        {
            "filename": "Приказ_по_производству_работ.docx",
            "title": "О назначении ответственного за производство строительных работ",
            "order_number": "17",
            "content": [
                "В целях обеспечения организации и контроля за качественным и безопасным выполнением строительных работ,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за производство строительных работ следующего сотрудника:",
                "   ФИО: Сидоров Сергей Сергеевич",
                "   Должность: Прораб",
                "2. Ответственному лицу обеспечить:",
                "   - контроль за выполнением работ в соответствии с проектной документацией и технологическими картами;",
                "   - соблюдение требований охраны труда и техники безопасности;",
                "   - координацию действий подрядчиков и рабочих;",
                "   - ведение необходимой технической и исполнительной документации.",
                "3. Контроль за исполнением настоящего приказа возложить на начальника строительного участка."
            ],
            "signatures": {"director": director_name, "responsible": "Сидоров С.С."}
        },
        {
            "filename": "Приказ_по_техническому_надзору.docx",
            "title": "О назначении ответственного за технический надзор ",
            "order_number": "18",
            "content": [
                "В целях контроля качества и безопасности строительно-монтажных работ,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за технический надзор на объекте:",
                "   ФИО: Кузнецова Марина Ивановна",
                "   Должность: Инженер технического надзора",
                "2. Ответственному лицу обеспечить:",
                "   - контроль соответствия выполняемых работ проектной документации;",
                "   - проверку соблюдения норм и стандартов качества;",
                "   - своевременное выявление и устранение нарушений;",
                "   - ведение отчетной документации по техническому надзору.",
                "3. Контроль за исполнением настоящего приказа возложить на начальника отдела технического контроля."
            ],
            "signatures": {"director": director_name, "responsible": "Кузнецова М.И."}
        },
        {
            "filename": "Приказ_по_огневым_работам.docx",
            "title": "О назначении ответственного за производство огневых работ",
            "order_number": "19",
            "content": [
                "В целях обеспечения безопасности при выполнении огневых работ на строительном объекте,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за производство огневых работ следующего сотрудника:",
                "   ФИО: Смирнов Алексей Викторович",
                "   Должность: Мастер строительных работ",
                "2. Ответственному лицу обеспечить:",
                "   - соблюдение требований пожарной безопасности при проведении огневых работ;",
                "   - организацию контроля за подготовкой рабочего места и средствами пожаротушения;",
                "   - проведение инструктажей с персоналом, выполняющим огневые работы;",
                "   - контроль за соблюдением правил охраны труда и техники безопасности.",
                "3. Контроль за исполнением настоящего приказа возложить на начальника строительного участка."
            ],
            "signatures": {"director": director_name, "responsible": "Смирнов А.В."}
        },
        {
            "filename": "Приказ_по_электробезопасности.docx",
            "title": "Об обеспечении электробезопасности",
            "order_number": "23",
            "content": [
                "В целях обеспечения электробезопасности ,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственного за электробезопасность следующего сотрудника:",
                "   ФИО: Сидоров Владимир Николаевич",
                "   Должность: Инженер по охране труда",
                "2. Провести инструктаж с сотрудниками по электробезопасности и обеспечить регулярное обновление знаний.",
                "3. Обеспечить техническое обслуживание и проверку электрооборудования.",
                "4. Контроль за исполнением настоящего приказа возложить на начальника отдела охраны труда."
            ],
            "signatures": {"director": director_name, "responsible": "Сидоров В.Н."}
        },
        {
            "filename": "Приказ_ответственного_на_объекте.docx",
            "title": "Об назначении ответственных лиц по следующим направлениям деятельности на объекте:",
            "order_number": "23",
            "content": [
                "В целях обеспечения организации и контроля за качественным и безопасным выполнением строительных работ ,",
                "ПРИКАЗЫВАЮ:",
                "1. Назначить ответственным за производство строительных работ следующего сотрудника:",
                "   ФИО: Сидоров Сергей Сергеевич",
                "   Должность: Прораб",
                "2. Ответственному лицу обеспечить:",
                "   - организацию безопасного выполнения высотных работ;",
                "   - контроль за соблюдением требований охраны труда и техники безопасности;",
                "   - проведение инструктажей и обучение персонала;",
                "   - контроль исправности средств индивидуальной защиты и страховочных систем.",
                "   - контроль за соблюдением требований пожарной безопасности;",
                "   - организацию проведения инструктажей по пожарной безопасности для работников;",
                "   - контроль исправности средств пожаротушения и эвакуационных путей;",
                "   - регулярную проверку и техническое обслуживание огнетушителей и другого оборудования;",
                "   - своевременное пополнение и замену средств пожаротушения;",
                "   - контроль за выполнением работ в соответствии с проектной документацией и технологическими картами;",
                "   - соблюдение требований охраны труда и техники безопасности;",
                "   - ведение необходимой технической и исполнительной документации.",
                "   - проверку соблюдения норм и стандартов качества;",
                "   - соблюдение требований пожарной безопасности при проведении огневых работ;",
                "   - организацию контроля за подготовкой рабочего места и средствами пожаротушения;",
                "   - проведение инструктажей с персоналом, выполняющим огневые работы;",
                "   - контроль за соблюдением правил охраны труда и техники безопасности.",
                "   - провести инструктаж с сотрудниками по электробезопасности и обеспечить регулярное обновление знаний.",
                "   - обеспечить техническое обслуживание и проверку электрооборудования.",
                "   - контроль за исполнением настоящего приказа возложить на начальника отдела охраны труда."
                "3. Контроль за исполнением настоящего приказа оставляю за собой."
            ],
            "signatures": {"director": director_name, "responsible": "Сидоров В.Н."}
        }
    ]

    for order in orders:
        create_order_doc(
            order["filename"],
            order["title"],
            city,
            date,
            order["order_number"],
            order["content"],
            order["signatures"]
        )

    print("Все приказы созданы, каждый — в отдельном файле, текст размещён на одной странице.")
